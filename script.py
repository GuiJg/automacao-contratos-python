import sys
from docx import Document
from docx.shared import Inches
from datetime import datetime
import locale
import os
import urllib.parse

# Configure a localidade para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Obter dados dos argumentos da linha de comando
contratante = sys.argv[1]
empresa = sys.argv[2]
cnpj = sys.argv[3]
rua = sys.argv[4]
bairro = sys.argv[5]
numero = sys.argv[6]
cidade = sys.argv[7]
estado = sys.argv[8]
cep = sys.argv[9]
orcamento = sys.argv[10]
vencimento = sys.argv[11]

# Criar um novo documento
documento = Document("contrato.docx")

# Dicionário de referências
referencias = {
    "XXXX": contratante,
    "PPPP": empresa,
    "FFFF": cnpj,
    "RRRR": rua,
    "BBBB": bairro,
    "MMMM": numero,
    "CCCC": cidade,
    "TTTT": estado,
    "CCCP": cep, 
    "OOOO": orcamento,
    "VVVV": vencimento,
    "DD": str(datetime.now().day),
    "MM": datetime.now().strftime("%B"),
    "AAAA": str(datetime.now().year),
}

# Substituir as referências no documento
for paragrafo in documento.paragraphs:
    for codigo in referencias:
        valor = referencias[codigo]
        paragrafo.text = paragrafo.text.replace(codigo, str(valor))

# Adicionar imagem ao documento
image_path = "logo.jpg"  # Substitua pelo caminho real da sua imagem
if os.path.exists(image_path):
    documento.add_picture(image_path, width=Inches(2))

# Salvar o documento com um nome específico
output_folder = "contratos"
os.makedirs(output_folder, exist_ok=True)
output_path = os.path.join(output_folder, f"Filiado - {contratante} - {cidade}.docx")
documento.save(output_path)

# Restaurar a localidade para a original (opcional)
locale.setlocale(locale.LC_TIME, '')

# Encode the output path for URL
encoded_output_path = urllib.parse.quote(output_path)
print(encoded_output_path)