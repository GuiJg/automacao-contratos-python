import sys
from docx import Document
from docx.shared import Inches
from datetime import datetime
import pandas as pd
import locale
import os

# Configure a localidade para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # ou 'pt_BR' dependendo do seu sistema

# Diretório de saída
output_dir = 'contratos'
os.makedirs(output_dir, exist_ok=True)

# Leitura do arquivo Excel
excel_file = sys.argv[1]
tabela = pd.read_excel(excel_file)

for linha in tabela.itertuples():
    # Criar um novo documento
    documento = Document("contrato.docx")

    # Obter dados da linha atual
    contratante = linha.contratante
    empresa = linha.empresa
    cpfj = linha.cpfj
    rua = linha.rua
    bairro = linha.bairro
    numero = linha.numero
    cidade = linha.cidade
    estado = linha.estado
    cep = linha.cep
    orcamento = linha.orcamento
    vencimento = linha.vencimento

    # Dicionário de referências
    referencias = {
        "XXXX": contratante,
        "PPPP": empresa,
        "FFFF": cpfj,
        "RRRR": rua,
        "BBBB": bairro,
        "MMMM": numero,
        "CCCC": cidade,
        "TTTT": estado,
        "CCCP": cep, 
        "OOOO": orcamento,
        "VVVV": vencimento,
        "DD": str(datetime.now().day),
        "MM": datetime.now().strftime("%B"),
        "AAAA": str(datetime.now().year)
    }

    # Substituir as referências no documento
    for paragrafo in documento.paragraphs:
        for codigo in referencias:
            valor = referencias[codigo]
            paragrafo.text = paragrafo.text.replace(codigo, str(valor))

    # Adicionar imagem ao documento
    image_path = "logo.jpg"  # Substitua pelo caminho real da sua imagem
    if os.path.exists(image_path):
        documento.add_picture(image_path, width=Inches(2))

    # Salvar o documento com um nome específico
    documento.save(os.path.join(output_dir, f"Filiação - {contratante} - {cidade}.docx"))
